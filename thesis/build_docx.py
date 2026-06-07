"""Convert thesis/Diplomatiki.tex into thesis/Diplomatiki.docx.

Lightweight LaTeX -> docx converter tailored to this template. Handles:
- \section / \subsection / \subsubsection -> Heading 1/2/3
- \begin{itemize|enumerate} ... \item -> bullet / numbered lists
- \begin{tabular} -> Word table
- $...$ and $$...$$ -> kept as raw text (math is not rendered, just preserved)
- \texttt{}, \textbf{}, \emph{}, \\, \newpage, comments -> normalised
- Front matter (title page, copyright, abstract, acknowledgements, TOC) is
  emitted as plain paragraphs.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent
TEX = ROOT / "Diplomatiki.tex"
DOCX = ROOT / "Diplomatiki.docx"

# Roots searched for \includegraphics{name} when name is not absolute.
# Mirrors the \graphicspath{} in the .tex preamble.
GRAPHICS_SEARCH_PATHS = [
    ROOT,
    ROOT / "figures",
    ROOT.parent / "databricks" / "docs",
]

FIGURE_DEFAULT_WIDTH = Inches(5.8)


# ---------- inline LaTeX cleanup ----------
INLINE_PATTERNS = [
    # typographic conversions: em/en dash and non-breaking space.
    # Must precede other text mangling. Order matters: --- before --.
    (re.compile(r"---"), "—"),
    (re.compile(r"--"), "–"),
    (re.compile(r"~"), " "),
    # repeated nested-brace flattening for common text macros (run twice to cover nesting)
    (re.compile(r"\\texttt\{([^{}]*)\}"), r"\1"),
    (re.compile(r"\\textbf\{([^{}]*)\}"), r"\1"),
    (re.compile(r"\\emph\{([^{}]*)\}"), r"\1"),
    (re.compile(r"\\textit\{([^{}]*)\}"), r"\1"),
    (re.compile(r"\\texttt\{([^{}]*)\}"), r"\1"),
    (re.compile(r"\\textbf\{([^{}]*)\}"), r"\1"),
    (re.compile(r"\\emph\{([^{}]*)\}"), r"\1"),
    (re.compile(r"\\textit\{([^{}]*)\}"), r"\1"),
    (re.compile(r"\\label\{[^{}]*\}"), ""),
    (re.compile(r"\\ref\{[^{}]*\}"), ""),
    (re.compile(r"\\hspace\{[^{}]*\}"), " "),
    (re.compile(r"\\vspace\{[^{}]*\}"), ""),
    (re.compile(r"\\setlength\{[^{}]*\}\{[^{}]*\}"), ""),
    (re.compile(r"\\parbox\s*(?:\[[a-z]\])?\s*\{[^{}]*\}\s*\{"), ""),
    (re.compile(r"\\includegraphics(?:\[[^\]]*\])?\{[^{}]*\}"), ""),
    (re.compile(r"\\caption\{([^{}]*)\}"), ""),
    # line break and space macros
    (re.compile(r"\\\\\s*"), "\n\n"),
    (re.compile(r"\\,"), " "),
    (re.compile(r"\\ (?=\S)"), " "),     # \<space> followed by char
    (re.compile(r"\\\s"), " "),          # \<whitespace>
    # escaped specials
    (re.compile(r"\\&"), "&"),
    (re.compile(r"\\%"), "%"),
    (re.compile(r"\\#"), "#"),
    (re.compile(r"\\_"), "_"),
    (re.compile(r"\\\$"), "$"),
    (re.compile(r"\\copyright\b"), "©"),
    # math/Greek shorthands
    (re.compile(r"\\to\b"), "→"),
    (re.compile(r"\\mid\b"), "|"),
    (re.compile(r"\\lambda\b"), "λ"),
    (re.compile(r"\\sigma\b"), "σ"),
    (re.compile(r"\\mu\b"), "μ"),
    (re.compile(r"\\Pi\b"), "Π"),
    # size, alignment, layout macros (strip)
    (re.compile(r"\\(?:Large|large|normalsize|small|scriptsize|tiny)\b"), ""),
    (re.compile(r"\\(?:centering|center|flushleft|flushright|noindent|newpage|clearpage|pagestyle\{[^{}]*\})\b"), ""),
    (re.compile(r"\\centering\b"), ""),
    (re.compile(r"\\tableofcontents\b"), "[Πίνακας Περιεχομένων]"),
    (re.compile(r"\\listoffigures\b"), "[Κατάλογος Εικόνων]"),
    (re.compile(r"\\listoftables\b"), "[Κατάλογος Πινάκων]"),
    (re.compile(r"\\appendix\b"), ""),
    (re.compile(r"\\date\{[^{}]*\}"), ""),
    (re.compile(r"\\title\{[^{}]*\}"), ""),
    (re.compile(r"\\begin\{center\}|\\end\{center\}"), ""),
    (re.compile(r"\\begin\{tabbing\}|\\end\{tabbing\}"), ""),
    (re.compile(r"\\begin\{large\}|\\end\{large\}"), ""),
    (re.compile(r"\\=\s"), " "),
    (re.compile(r"\\>\s"), " "),
    # math delimiters: drop $$ ... $$ and $ ... $ wrappers, keep contents
    (re.compile(r"\$\$(.+?)\$\$", re.DOTALL), r"[\1]"),
    (re.compile(r"\$([^$]+)\$"), r"\1"),
    # flatten leftover single-brace groups around short text (run twice for nesting)
    (re.compile(r"\{([^{}\\]{1,40})\}"), r"\1"),
    (re.compile(r"\{([^{}\\]{1,40})\}"), r"\1"),
    # drop empty braces and any remaining unknown \command tokens
    (re.compile(r"\{\}"), ""),
    (re.compile(r"\\[a-zA-Z]+\*?"), ""),
    # cleanup orphan braces
    (re.compile(r"(?<!\\)[{}]"), ""),
    # unescape LaTeX-escaped braces \{ and \} -> { and }
    (re.compile(r"\\\{"), "{"),
    (re.compile(r"\\\}"), "}"),
    # whitespace cleanup
    (re.compile(r"[ \t]+"), " "),
    (re.compile(r"\n[ \t]+"), "\n"),
    (re.compile(r"\n{3,}"), "\n\n"),
]


_NOISE_LINE_RE = re.compile(r"^[\s\\{}\.]*$")


def clean_inline(text: str) -> str:
    for pattern, repl in INLINE_PATTERNS:
        text = pattern.sub(repl, text)
    # drop lines that became pure punctuation/noise; keep blank lines so
    # paragraph breaks (\n\n) survive into flush_text()
    kept: list[str] = []
    for ln in text.split("\n"):
        if ln.strip() == "":
            kept.append("")
        elif _NOISE_LINE_RE.match(ln):
            continue
        else:
            kept.append(ln)
    out = "\n".join(kept)
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out.strip()


def parse_tabular(body: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for raw in body.split("\\\\"):
        # strip \hline tokens first; only skip if nothing else remains
        raw = re.sub(r"\\hline\b", "", raw).strip()
        if not raw:
            continue
        cells = [clean_inline(c) for c in raw.split("&")]
        rows.append(cells)
    return rows


# ---------- figure preprocessing ----------
_FIGURE_RE = re.compile(
    r"\\begin\{figure\}(?:\[[^\]]*\])?(.*?)\\end\{figure\}",
    re.DOTALL,
)
_INCLUDEGRAPHICS_RE = re.compile(
    r"\\includegraphics(?:\[[^\]]*\])?\{([^{}]+)\}"
)
# Caption body may contain one level of nested {...} (e.g. \texttt{x}); allow it.
_CAPTION_RE = re.compile(
    r"\\caption\{((?:[^{}]|\{[^{}]*\})*)\}"
)


def _resolve_image(name: str) -> Path | None:
    name = name.strip()
    if not name:
        return None
    p = Path(name)
    if p.is_absolute() and p.exists():
        return p
    for root in GRAPHICS_SEARCH_PATHS:
        cand = (root / name).resolve()
        if cand.exists():
            return cand
    # try with common extensions appended
    for ext in (".png", ".jpg", ".jpeg", ".pdf", ".eps"):
        for root in GRAPHICS_SEARCH_PATHS:
            cand = (root / (name + ext)).resolve()
            if cand.exists():
                return cand
    return None


def preprocess_figures(body: str) -> str:
    """Replace each figure environment with a single-line sentinel:
        [[FIGURE|<absolute path or empty>|<raw image name>|<caption>]]
    Sentinels survive the rest of the cleanup unchanged because they contain
    no backslashes or braces.
    """
    def _sub(match: re.Match[str]) -> str:
        inner = match.group(1)
        # \includegraphics may be commented out (line starting with %); strip
        # leading % only when looking for the path
        inc_match = _INCLUDEGRAPHICS_RE.search(re.sub(r"(?m)^\s*%\s*", "", inner))
        cap_match = _CAPTION_RE.search(inner)
        raw_name = inc_match.group(1) if inc_match else ""
        caption = cap_match.group(1) if cap_match else ""
        resolved = _resolve_image(raw_name) if raw_name else None
        path_str = str(resolved) if resolved else ""
        # Strip pipe/newline from caption to keep sentinel single-line
        caption = caption.replace("|", "/").replace("\n", " ").strip()
        return f"\n\n[[FIGURE|{path_str}|{raw_name}|{caption}]]\n\n"

    return _FIGURE_RE.sub(_sub, body)


_FIGURE_SENTINEL_RE = re.compile(
    r"\[\[FIGURE\|([^|\]]*)\|([^|\]]*)\|([^\]]*)\]\]"
)


_BIBLIOGRAPHY_RE = re.compile(
    r"\\begin\{thebibliography\}\{[^{}]*\}(.*?)\\end\{thebibliography\}",
    re.DOTALL,
)


# Author-year display for \cite{Key} expansion in the DOCX renderer.
# Keys must match the \bibitem labels in Diplomatiki.tex.
CITE_LABELS: dict[str, str] = {
    "BaroneAdesi2008": "Barone-Adesi, Engle και Mancini, 2008",
    "Bates2000": "Bates, 2000",
    "BlackScholes1973": "Black και Scholes, 1973",
    "Coleman": "Coleman κ.ά.",
    "Damji2020": "Damji κ.ά., 2020",
    "DuanSimonato1998": "Duan και Simonato, 1998",
    "Engle1982": "Engle, 1982",
    "Gavrilov": "Gavrilov κ.ά.",
    "Heston1993": "Heston, 1993",
    "Jiang2004": "Jiang, 2004",
    "MandelbrotFC1997": "Mandelbrot, Fisher και Calvet, 1997",
    "Merton1973": "Merton, 1973",
    "Merton1976": "Merton, 1976",
    "Moldovan": "Moldovan κ.ά.",
    "Mulvey1992": "Mulvey και Vladimirou, 1992",
    "Nandi1998": "Nandi, 1998",
    "PaparoditisPolitis2002": "Paparoditis και Politis, 2002",
    "PolitisRomano1994": "Politis και Romano, 1994",
    "RockafellarWets1991": "Rockafellar και Wets, 1991",
    "Shanker1996": "Shanker, Hu και Hung, 1996",
    "ZhangBSDE2011": "Zhang κ.ά., 2011",
    "Zhou2011": "Zhou, Lei και Ye, 2011",
}


_CITE_RE = re.compile(r"\\cite\{([^{}]+)\}")
_CITEYEAR_RE = re.compile(r"\\citeyear\{([^{}]+)\}")


def _label_year(label: str) -> str:
    m = re.search(r"(\d{4})", label)
    return m.group(1) if m else "χ.χ."


def expand_cites(body: str) -> str:
    """Replace \\cite{KeyA, KeyB} with (Label A; Label B) and \\citeyear{Key}
    with (Year), before the generic LaTeX cleanup strips unknown commands."""
    def _sub_cite(match: re.Match[str]) -> str:
        keys = [k.strip() for k in match.group(1).split(",") if k.strip()]
        labels = [CITE_LABELS.get(k, k) for k in keys]
        return "(" + "· ".join(labels) + ")"

    def _sub_year(match: re.Match[str]) -> str:
        keys = [k.strip() for k in match.group(1).split(",") if k.strip()]
        years = [_label_year(CITE_LABELS.get(k, k)) for k in keys]
        return "(" + ", ".join(years) + ")"

    body = _CITEYEAR_RE.sub(_sub_year, body)
    body = _CITE_RE.sub(_sub_cite, body)
    return body


def preprocess_bibliography(body: str) -> str:
    """Rewrite \\begin{thebibliography}{99} ... \\end{thebibliography} into a
    \\section{Βιβλιογραφία} followed by one paragraph per \\bibitem entry, so
    the existing tokenizer renders it without special-casing."""
    def _sub(match: re.Match[str]) -> str:
        inner = match.group(1)
        items = re.split(r"\\bibitem\{[^{}]*\}", inner)
        parts = []
        for p in items:
            if not p.strip():
                continue
            p = re.sub(r"\s+", " ", p).strip()
            p = p.replace("~", " ")
            p = p.replace("--", "–")
            p = re.sub(r"\\=a", "ā", p)
            parts.append(p)
        out = ["\\section{Βιβλιογραφία}"]
        out.extend(parts)
        return "\n\n" + "\n\n".join(out) + "\n\n"

    return _BIBLIOGRAPHY_RE.sub(_sub, body)


def _add_figure(doc: Document, abs_path: str, raw_name: str, caption: str) -> None:
    caption_clean = clean_inline(caption) if caption else ""
    if abs_path and Path(abs_path).exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        try:
            run.add_picture(abs_path, width=FIGURE_DEFAULT_WIDTH)
        except Exception as exc:  # noqa: BLE001
            para.add_run(f"[Σφάλμα φόρτωσης εικόνας: {raw_name} ({exc})]")
    else:
        placeholder = doc.add_paragraph()
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        msg = (
            f"[ΕΙΚΟΝΑ ΠΡΟΣ ΠΡΟΣΘΗΚΗ — δεν εντοπίστηκε αρχείο "
            f"({raw_name or 'χωρίς διαδρομή'})]"
        )
        run = placeholder.add_run(msg)
        run.bold = True
    if caption_clean:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(f"Εικόνα: {caption_clean}")
        cap_run.italic = True


def main() -> None:
    src = TEX.read_text(encoding="utf-8")

    # strip LaTeX comments (lines starting with %, and trailing % on a line)
    src = re.sub(r"(?<!\\)%.*", "", src)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cut everything between \begin{document} and \end{document}
    match = re.search(r"\\begin\{document\}(.*)\\end\{document\}", src, re.DOTALL)
    body = match.group(1) if match else src

    # Convert figure environments into [[FIGURE|...]] sentinels before any
    # other processing so the picture is rendered exactly once, in place.
    body = preprocess_figures(body)

    # Expand the thebibliography environment into a plain section + paragraphs.
    body = preprocess_bibliography(body)

    # Expand \cite{Key} macros to author-year strings before generic cleanup.
    body = expand_cites(body)

    # tokenize section by section
    # split on section/subsection/subsubsection while keeping the markers.
    # Heading bodies may contain one level of nested braces (e.g. \texttt{x}).
    _H = r"(?:[^{}]|\{[^{}]*\})+"
    # tabular column specs may contain nested braces (e.g. p{4.5cm}).
    _CS = r"(?:[^{}]|\{[^{}]*\})*"
    token_re = re.compile(
        r"(\\section\{" + _H + r"\}|\\subsection\{" + _H + r"\}|\\subsubsection\{" + _H + r"\}|"
        r"\\begin\{itemize\}|\\end\{itemize\}|\\begin\{enumerate\}|\\end\{enumerate\}|"
        r"\\begin\{table\}\[?[^\]]*\]?|\\end\{table\}|"
        r"\\begin\{tabular\}\{" + _CS + r"\}|\\end\{tabular\}|"
        r"\\item)"
    )

    pos = 0
    in_list = None  # "bullet" | "number" | None
    in_tabular = False
    tabular_buf: list[str] = []

    def flush_text(txt: str) -> None:
        """Render a free-text chunk, splitting out FIGURE sentinels so they
        become real pictures instead of literal text."""
        if not txt:
            return
        last = 0
        for fm in _FIGURE_SENTINEL_RE.finditer(txt):
            before = txt[last:fm.start()]
            cleaned = clean_inline(before)
            if cleaned:
                for para in re.split(r"\n\s*\n", cleaned):
                    para = para.strip()
                    if para:
                        doc.add_paragraph(para)
            _add_figure(doc, fm.group(1), fm.group(2), fm.group(3))
            last = fm.end()
        tail = clean_inline(txt[last:])
        if tail:
            for para in re.split(r"\n\s*\n", tail):
                para = para.strip()
                if para:
                    doc.add_paragraph(para)

    for m in token_re.finditer(body):
        chunk = body[pos:m.start()]
        token = m.group(1)
        pos = m.end()

        if in_tabular:
            tabular_buf.append(chunk)
            if token.startswith("\\end{tabular}"):
                rows = parse_tabular("".join(tabular_buf))
                if rows:
                    ncols = max(len(r) for r in rows)
                    table = doc.add_table(rows=len(rows), cols=ncols)
                    table.style = "Light Grid Accent 1"
                    for i, r in enumerate(rows):
                        for j, c in enumerate(r):
                            table.cell(i, j).text = c
                tabular_buf = []
                in_tabular = False
            continue

        # normal chunk
        if chunk.strip():
            if in_list:
                # \item handled separately; remainder is item text
                txt = clean_inline(chunk)
                if txt:
                    doc.add_paragraph(
                        txt,
                        style="List Bullet" if in_list == "bullet" else "List Number",
                    )
            else:
                flush_text(chunk)

        if token.startswith("\\section{"):
            heading = clean_inline(token[len("\\section{"):-1])
            doc.add_heading(heading, level=1)
        elif token.startswith("\\subsection{"):
            heading = clean_inline(token[len("\\subsection{"):-1])
            doc.add_heading(heading, level=2)
        elif token.startswith("\\subsubsection{"):
            heading = clean_inline(token[len("\\subsubsection{"):-1])
            doc.add_heading(heading, level=3)
        elif token == "\\begin{itemize}":
            in_list = "bullet"
        elif token == "\\end{itemize}":
            in_list = None
        elif token == "\\begin{enumerate}":
            in_list = "number"
        elif token == "\\end{enumerate}":
            in_list = None
        elif token == "\\item":
            pass  # next chunk is item text
        elif token.startswith("\\begin{tabular}"):
            in_tabular = True
            tabular_buf = []
        elif token.startswith("\\begin{table}"):
            pass
        elif token == "\\end{table}":
            pass

    # tail
    if pos < len(body):
        flush_text(body[pos:])

    doc.save(DOCX)
    print(f"Wrote {DOCX}")


if __name__ == "__main__":
    main()
